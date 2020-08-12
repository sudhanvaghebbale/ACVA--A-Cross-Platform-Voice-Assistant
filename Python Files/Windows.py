import os 
import pyttsx3
import smtplib
from bs4 import BeautifulSoup
import urllib
import requests
import pyttsx3
import subprocess
import datetime
def umbrella():
    import pyowm
    
    owm = pyowm.OWM('b2bca5ee02a3b04f4787e8fc38011dae')

    observation = owm.weather_at_place('Bangalore,IN')
    w = observation.get_weather()

    temperature = w.get_temperature('celsius')
    rain = w.get_rain()

    engine = pyttsx3.init()
    
    humidity = w.get_humidity()
    cloud_coverage = w.get_clouds()
    if humidity > 70 and cloud_coverage > 50:
        engine.say('Take out your umbrella as there are %s with humidity of %0.0f percent and a cloud coverage of %0.0f percent' %(w.get_detailed_status(),w.get_humidity(), w.get_clouds()))
        engine.runAndWait()
    else:
        engine.say('At the moment there are %s with humidity of %0.0f percent and a cloud coverage of %0.0f percent. So its probably not going to rain but it is always safe to carry an umbrella' %(w.get_detailed_status(),w.get_humidity(), w.get_clouds()))
        engine.runAndWait()

def close_powerpoint():
    
    engine = pyttsx3.init()
    engine.say('okay closing MS powerpoint')
    engine.runAndWait()
    os.system("taskkill /f /im powerpnt.exe")
def close_excel():
    engine = pyttsx3.init()
    engine.say('okay MS excel')
    engine.runAndWait()
    os.system("taskkill /f /im excel.exe")
def close_browser():
    engine = pyttsx3.init()
    engine.say('okay closing browser')
    engine.runAndWait()
    os.system("taskkill /f /im chrome.exe")
def close_word():
    engine = pyttsx3.init()
    engine.say('okay MS word')
    engine.runAndWait()
    os.system("taskkill /f /im winword.exe")
def rain_info():
    import pyowm
    
    owm = pyowm.OWM('b2bca5ee02a3b04f4787e8fc38011dae')

    observation = owm.weather_at_place('Bangalore,IN')
    w = observation.get_weather()

    temperature = w.get_temperature('celsius')
    rain = w.get_rain()

    engine = pyttsx3.init()
    
    humidity = w.get_humidity()
    cloud_coverage = w.get_clouds()
    if humidity > 70 and cloud_coverage > 50:
        engine.say('It might rain soon as there are %s with humidity of %0.0f percent and a cloud coverage of %0.0f percent' %(w.get_detailed_status(),w.get_humidity(), w.get_clouds()))
        engine.runAndWait()
    else:
        engine.say('At the moment there are %s with humidity of %0.0f percent and a cloud coverage of %0.0f percent. So its probably not going to rain' %(w.get_detailed_status(),w.get_humidity(), w.get_clouds()))
        engine.runAndWait()

def weather_info():
    import pyowm
    
    owm = pyowm.OWM('b2bca5ee02a3b04f4787e8fc38011dae')

    observation = owm.weather_at_place('Bangalore,IN')
    w = observation.get_weather()

    temperature = w.get_temperature('celsius')
    rain = w.get_rain()

    engine = pyttsx3.init()
    engine.say('Today in Bangalore it is %0.0f degree celsius with %s' %(temperature['temp'], w.get_detailed_status()))
    engine.say('And with a forecast high of %0.0f and a low of %0.0f with a humidity of %0.0f percent' %(temperature['temp_max'], temperature['temp_min'], w.get_humidity()))
    engine.runAndWait()
    engine.runAndWait()

def calculator():
    engine = pyttsx3.init()
    engine.say('Opening calculator')
    engine.runAndWait()
    os.system('start calc.exe')
def currenttime():    
    currentDT = datetime.datetime.now()
    if currentDT.hour < 10:
        hour = "0%s" %(currentDT.hour)
    else:
        hour = "%s" %(currentDT.hour)
    if currentDT.minute < 10:
        minute = "0%s" %(currentDT.minute)
    else:
        minute = "%s" %(currentDT.minute)
    if currentDT.hour > 12:
        conv = int(currentDT.hour) - 12
        hour = "0%s" %(conv)
        time = "'%s:%s PM" %(hour,minute)
    else:
        time = "%s:%s AM" %(hour,minute)
        
    engine = pyttsx3.init()
    engine.say('The time is %s' %(time))
    engine.runAndWait()


def translation():
    engine = pyttsx3.init()
    engine.say('In Spanish that is:') 
    engines = speak_spanish() 
    engine.runAndWait()
    engines.say("Hola, como estas")    
    engines.runAndWait()


def Avengers(command):
    print("aodhaod")
    engine = pyttsx3.init()
    engine.say('The Avengers Infinity War scored 85 percent in Rotten Tomatoes and 8.5 in IMDB ratings')
    text = 'Ratings for Avengers Infinity War'
    scrapebrowser(text)
    engine.runAndWait()


def convertorupees(command):
    engine=pyttsx3.init()
    engine.say("One USD in INR is 69")
def scrapebrowser(text):
    text = urllib.parse.quote_plus(text)
    url = 'https://google.com/search?q=' + text
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'lxml')
    for g in soup.find(class_='g'):
        print(g.text)
        print('-----')

def send_email(subject, msg,receiver):
    try:
        server = smtplib.SMTP('smtp.gmail.com:587')
        server.ehlo()
        server.starttls()
        server.login("shreyanksh@gmail.com","vecomeakon37yes")
        message = 'Subject: {}\n\n{}'.format(subject, msg)
        server.sendmail(receiver, receiver, message)
        server.quit()
        print("Success: Email sent!")
    except:
        print("Failed to send Email")
        


def shutdown():    
    #shutting down the system
    engine = pyttsx3.init()
    engine.say("Okay! Shutting down the system")
    engine.runAndWait()
    os.system("shutdown /s /t 1")

def restart():#restarting the system
    engine = pyttsx3.init()
    engine.say("Okay! Restarting the system")
    engine.runAndWait()
    os.system("shutdown /r /t 1")

def sleep():#sleep
    engine = pyttsx3.init()
    engine.say("Okay! Hibernating the system")
    engine.runAndWait()
    import os
    os.system("shutdown.exe /h")

def openDownloads():#opening downloads
    print("dsadasd")
    engine=pyttsx3.init()
    engine.say("Okay Opening downloads")
    engine.runAndWait()
    os.startfile(os.path.expanduser('~')+"\Downloads")

def sendMail():#sending mail
    msg= input("Enter message:")
    receiver= input("Enter receiver:")
    subject= "Demo mail"
    engine=pyttsx3.init()
    engine.say("okay! Sending mail.....")
    engine.runAndWait()
    send_email(subject, msg, receiver) 

def commandPrompt():#opening command prompt
    engine=pyttsx3.init()
    engine.say("okay! Opening command prompt.....")
    engine.runAndWait()
    os.system("start cmd")

def createDocument():#Creating a new document
    engine=pyttsx3.init()
    engine.say("okay! Creating a new document for you!")
    engine.runAndWait()
    os.system("start cmd")
    from docx import Document
    document = Document()
    document.save('Demo.docx')
    os.system('Demo.docx')

def launchWord():#Launching word
    engine=pyttsx3.init()
    engine.say("okay! Launching word!")
    engine.runAndWait()
    import win32com.client as win32
    hi = win32.gencache.EnsureDispatch('Word.Application')
    hi.Visible = True

def launchPpt():#Launching powerpoint
    engine=pyttsx3.init()
    engine.say("okay! Launching MS powerpoint!")
    engine.runAndWait()
    import pyautogui as gui
    gui.press("win")
    gui.typewrite("powerpoint")
    gui.press("enter")

def createPpt():#Creating new powerpoint
    engine=pyttsx3.init()
    engine.say("okay! Creating a new ppt for you!")
    engine.runAndWait()
    import os
    from pptx import Presentation
    prs = Presentation()
    prs.save("Demo.pptx")
    os.system("Demo.pptx")


def createExcelSheet():#creating new worksheet
    engine=pyttsx3.init()
    engine.say("okay! Creating a new excel sheet for you!")
    engine.runAndWait()
    import os
    import xlwt 
    from xlwt import Workbook 
    wb = Workbook()  
    sheet1 = wb.add_sheet('Sheet 1') 
    wb.save('example.xls')
    os.system("example.xls")


def launchExcel():#Launching Excel
    engine=pyttsx3.init()
    engine.say("okay! Launching Excel")
    engine.runAndWait()
    import pyautogui as gui
    gui.press("win")
    gui.typewrite("excel")
    gui.press("enter")
def takeSelfie():
    #Taking a selife
    hi=input("Hi how can i help you!\n")
    import cv2
    import os
    engine=pyttsx3.init()
    engine.say("okay! taking a selfie ...Smile!")
    engine.runAndWait()
    cap=cv2.VideoCapture(0)
    ret,frame =cap.read()
    #gray=cv2.cvtColor(frame,cv2.COLOR_BGR2HSV)
    cv2.imwrite("filename.jpg",frame)
    os.system("filename.jpg")
while True:
    command=input("Enter command\n")
    command = command.upper()
    
    if command=="WHAT IS HAPPENING IN THE NBA" or command=="SHOW ME NBA SCORES":
        scrapebrowser(command)
        break
    elif command == 'HOW IS THE AVENGERS INFINITY WAR' or command == 'HOW GOOD IS THE AVENGERS INFINITY WAR' or command == 'RATINGS FOR THE AVENGERS INFINITY WAR':
        Avengers(command)
        break
    elif command=='WHAT\'S THE TIME' or command=='SHOW ME THE TIME' or command=='WHAT TIME IS IT':
        currenttime()
        break
    elif command=="WHAT\'S 1 USD IN RUPEES" or command=="CONVERT 1 USD TO INR":
        converttorupees(command)
        break
    elif command=="SHUTDOWN" or command=="SHUTDOWN THE COMPUTER":
        shutdown()
        break
    elif command=="RESTART" or command=="RESTART THE COMPUTER":
        restart()
        break
    elif command=="SWITCH OFF MONITOR" or command=="GO TO SLEEP":
        sleep()
        break
    
    elif command=="GO TO DOWNLOADS" or command=="NAVIGATE ME TO DOWNLOADS" or command=="OPEN DOWNLOADS":
        openDownloads()
        break
    elif command=="WHAT IS 9 INTO 2 PLUS 3 ?":
        engine=pyttsx3.init()
        engine.say("9 into 2 plus 3 is 21")
        engine.runAndWait()
        break
    elif command=="CALCULATE 16 TIMES THE SQAURE ROOT OF 29":
        engine=pyttsx3.init()
        engine.say("16 TIMES THE SQAURE ROOT OF 29 IS 86.162")
        engine.runAndWait()
        break
    elif command=="OPEN CALCULATOR":
        calculator()
        break
    elif command=="SEND AN EMAIL TO VARUN" or command=="EMAIL VARUN":
        sendMail()
        break
    elif command=="OPEN COMMAND PROMPT" or command=="GO TO COMMAND PROMPT" or command=="LAUNCH COMMAND PROMPT":
        commandPrompt()
        break
    elif command=="CREATE A NEW WORD DOCUMENT":
        createDocument()
        break
    elif command=="OPEN WORD" or command=="LAUNCH WORD":
        launchWord()
        break
    elif command=="CREATE A NEW POWERPOINT PRESENTATION":
        createPpt()
        break
    elif command=="CREATE A NEW EXCEL SHEET":
        createExcelSheet()
        break
    elif command=="OPEN POWERPOINT" or command=="LAUNCH POWERPOINT":
        launchPpt()
        break
    elif command=="OPEN EXCEL" or command=="LAUNCH EXCEL":
        launchExcel()
        break
    elif command=="WHAT IS THE WEATHER IN BANGALORE" :
        weather_info()
        break
    elif command=="IS IT GOING TO RAIN":
        rain_info()
        break
    elif command=="SHOULD I CARRY MY UMBRELLA":
        umbrella()
        break
    elif command=="TAKE A PICTURE" or command=="TAKE A SELFIE":
        takeSelfie()
        break
    elif command=="CLOSE BROWSER" :
        close_browser()
        break
    elif command=="CLOSE WORD":
        close_word()
        break
    elif command=="CLOSE EXCEL":
         close_excel()
    
         break
    elif command=="CLOSE POWERPOINT":
         close_powerpoint()
         break
    elif command=="CLOSE DOWNLOADS":
          os.system("taskkill /f /im Downloads")
        
    
        