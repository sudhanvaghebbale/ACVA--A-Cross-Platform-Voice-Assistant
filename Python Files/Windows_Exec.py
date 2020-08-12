import os 
import pyttsx3
import smtplib
from bs4 import BeautifulSoup
import urllib
import requests
import pyttsx3
import subprocess
import datetime
import pyautogui as gui
import win32com.client as win32
from docx import Document
import xlwt 
from xlwt import Workbook 
import pyowm
import cv2
from pptx import Presentation
import webbrowser
engine = pyttsx3.init()
rate = engine.getProperty('rate')
engine.setProperty('rate', 125)
voices = engine.getProperty('voices')
engine.setProperty('voice', voices[1].id)

def umbrella():
   
    
    owm = pyowm.OWM('b2bca5ee02a3b04f4787e8fc38011dae')

    observation = owm.weather_at_place('Bangalore,IN')
    w = observation.get_weather()

    temperature = w.get_temperature('celsius')
    rain = w.get_rain()

    engine = pyttsx3.init()
    
    humidity = w.get_humidity()
    cloud_coverage = w.get_clouds()
    if humidity > 70 and cloud_coverage > 50:
        engine.say('Take out your umbrella as there are %s with humidity of %0.0f percent and a cloud coverage of %0.0f percent' %(w.get_detailed_status(),w.get_humidity(), w.get_clouds()))
        engine.runAndWait()
    else:
        engine.say('At the moment there are %s with humidity of %0.0f percent and a cloud coverage of %0.0f percent. So its probably not going to rain but it is always safe to carry an umbrella' %(w.get_detailed_status(),w.get_humidity(), w.get_clouds()))
        engine.runAndWait()

def close_powerpoint():
    engine = pyttsx3.init()
    engine.say('okay closing MS powerpoint')
    engine.runAndWait()
    os.system("taskkill /f /im powerpnt.exe")
    closeppt = "Okay closing MS Powerpoint"
    return closeppt
def close_excel():
    engine = pyttsx3.init()
    engine.say('okay closing MS excel')
    engine.runAndWait()
    os.system("taskkill /f /im excel.exe")
    closeexcel = "Okay closing MS Excel"
    return closeexcel
def close_browser():
    engine = pyttsx3.init()
    engine.say('okay closing browser')
    engine.runAndWait()
    os.system("taskkill /f /im chrome.exe")
    closebrowser = "Okay closing Browser"
    return closebrowser
def close_word():
    engine = pyttsx3.init()
    engine.say('okay closing MS word')
    engine.runAndWait()
    os.system("taskkill /f /im winword.exe")
    closeword = "Okay closing MS Word"
    return closeword
def rain_info():
    import pyowm
    
    owm = pyowm.OWM('b2bca5ee02a3b04f4787e8fc38011dae')

    observation = owm.weather_at_place('Bangalore,IN')
    w = observation.get_weather()

    temperature = w.get_temperature('celsius')
    rain = w.get_rain()

    engine = pyttsx3.init()
    
    humidity = w.get_humidity()
    cloud_coverage = w.get_clouds()
    if humidity > 70 and cloud_coverage > 50:
        engine.say('It might rain soon as there are %s with humidity of %0.0f percent and a cloud coverage of %0.0f percent' %(w.get_detailed_status(),w.get_humidity(), w.get_clouds()))
        engine.runAndWait()
    else:
        engine.say('At the moment there are %s with humidity of %0.0f percent and a cloud coverage of %0.0f percent. So its probably not going to rain' %(w.get_detailed_status(),w.get_humidity(), w.get_clouds()))
        engine.runAndWait()

def weather_information():
    import pyowm
    
    owm = pyowm.OWM('b2bca5ee02a3b04f4787e8fc38011dae')

    observation = owm.weather_at_place('Bangalore,IN')
    w = observation.get_weather()

    temperature = w.get_temperature('celsius')
    rain = w.get_rain()

    status = w.get_detailed_status()
    max = temperature['temp_max']
    min = temperature['temp_min']
    humid = w.get_humidity()
    return status, max,min,humid

def rain_information():
    import pyowm
    
    owm = pyowm.OWM('b2bca5ee02a3b04f4787e8fc38011dae')

    observation = owm.weather_at_place('Bangalore,IN')
    w = observation.get_weather()

    temperature = w.get_temperature('celsius')
    rain = w.get_rain()

    humidity = w.get_humidity()
    cloud_coverage = w.get_clouds()
    if humidity < 50 and cloud_coverage > 50:
        
        status = w.get_detailed_status()
        humid = w.get_humidity()
        clouds = w.get_clouds()
        text = 'Yes'
        return status,humid,clouds,text
    else:

        status = w.get_detailed_status()
        humid = w.get_humidity()
        clouds = w.get_clouds()
        text = 'No'
        return status,humid,clouds,text

def current_time():    
    currentDT = datetime.datetime.now()
    if currentDT.hour < 10:
        hour = "0%s" %(currentDT.hour)
    else:
        hour = "%s" %(currentDT.hour)
    if currentDT.minute < 10:
        minute = "0%s" %(currentDT.minute)
    else:
        minute = "%s" %(currentDT.minute)
    if currentDT.hour > 12:
        conv = int(currentDT.hour) - 12
        hour = "0%s" %(conv)
        time = "'%s:%s PM" %(hour,minute)
    else:
        time = "%s:%s AM" %(hour,minute)
        
    return time

def weather_info():
    import pyowm
    
    owm = pyowm.OWM('b2bca5ee02a3b04f4787e8fc38011dae')

    observation = owm.weather_at_place('Bangalore,IN')
    w = observation.get_weather()

    temperature = w.get_temperature('celsius')
    rain = w.get_rain()

    engine = pyttsx3.init()
    engine.say('Today in Bangalore it is %0.0f degree celsius with %s' %(temperature['temp'], w.get_detailed_status()))
    engine.say('And with a forecast high of %0.0f and a low of %0.0f with a humidity of %0.0f percent' %(temperature['temp_max'], temperature['temp_min'], w.get_humidity()))
    engine.runAndWait()
    engine.runAndWait()

def calculator():
    engine = pyttsx3.init()
    engine.say('Opening calculator')
    engine.runAndWait()
    os.system('start calc.exe')
    calculator = "Opening Calculator"
    return calculator
def currenttime():    
    currentDT = datetime.datetime.now()
    if currentDT.hour < 10:
        hour = "0%s" %(currentDT.hour)
    else:
        hour = "%s" %(currentDT.hour)
    if currentDT.minute < 10:
        minute = "0%s" %(currentDT.minute)
    else:
        minute = "%s" %(currentDT.minute)
    if currentDT.hour > 12:
        conv = int(currentDT.hour) - 12
        hour = "0%s" %(conv)
        time = "'%s:%s PM" %(hour,minute)
    else:
        time = "%s:%s AM" %(hour,minute)
        
    engine = pyttsx3.init()
    engine.say('The time is %s' %(time))
    engine.runAndWait()
def speak_spanish():
    engine = pyttsx3.init()
    voices = engine.getProperty('voices')
    for voice in voices:
        if voice.languages[0] == u'es_ES':
            engine.setProperty('voice', voice.id)
            break

def translation():
    engine = pyttsx3.init()
    engine.say('In Spanish that is:') 
    engine.runAndWait()
    engine.say("Hola, como estas")    
    engine.runAndWait()
    spanish = "In Spanish that is: Hola, como estas"
    return spanish

def Avengers(command):
    engine = pyttsx3.init()
    engine.say('The Avengers Infinity War scored 85 percent in Rotten Tomatoes and 8.5 in IMDB ratings')
    text = 'Ratings for Avengers Infinity War'
    #scrapebrowser(text)
    engine.runAndWait()
    avengers = "The Avengers Infinity War scored 85 percent in Rotten Tomatoes and 8.5 in IMDB ratings"
    return avengers

def convertorupees(command):
    
    engine=pyttsx3.init()
    engine.say("One USD in INR is 69")
    rupees = "One USD in INR is 69"
    return rupees
def scrapebrowser(text):
    webbrowser.open("https://www.basketball-reference.com/boxscores/")
    engine=pyttsx3.init()
    engine.say("Here are the latest updates in the NBA")
    engine.runAndWait()
    rupees = "Here are the latest updates in the NBA"
    return rupees

def send_email(subject, msg,receiver):
    try:
        server = smtplib.SMTP('smtp.gmail.com:587')
        server.ehlo()
        server.starttls()
        server.login("sudhanva711@gmail.com","sudhanva711")
        message = 'Subject: {}\n\n{}'.format(subject, msg)
        server.sendmail(receiver, receiver, message)
        server.quit()
        print("Success: Email sent!")
        return "Email sent!"
    except:
        print("Failed to send Email")
        return "Failed to send email"
        


def shutdown():    
    #shutting down the system
    engine = pyttsx3.init()
    engine.say("Okay! Shutting down the system")
    engine.runAndWait()
    os.system("shutdown /s /t 1")
    shutdown = "Okay! Shutting down the system"
    return shutdown

def restart():#restarting the system
    engine = pyttsx3.init()
    engine.say("Okay! Restarting the system")
    engine.runAndWait()
    os.system("shutdown /r /t 1")
    restart = "Okay! Restarting the system"
    return restart

def sleep():#sleep
    engine = pyttsx3.init()
    engine.say("Okay! Hibernating the system")
    engine.runAndWait()
    import os
    os.system("shutdown.exe /h")
    return "Okay! Hibernating the system"

def openDownloads():#opening downloads
    engine=pyttsx3.init()
    engine.say("Okay Opening downloads")
    engine.runAndWait()
    os.startfile(os.path.expanduser('~')+"\Downloads")
    downloads = "Okay Opening Downloads"
    return downloads

def sendMail():#sending mail
    msg= input("Enter message:")
    receiver= input("Enter receiver:")
    subject= "Demo mail"
    engine=pyttsx3.init()
    engine.say("okay! Sending mail.....")
    engine.runAndWait()
    send_email(subject, msg, receiver) 

def commandPrompt():#opening command prompt
    engine=pyttsx3.init()
    engine.say("okay! Opening command prompt.....")
    engine.runAndWait()
    os.system("start cmd")
    return "Okay! Opening Command Prompt!"

def createDocument():#Creating a new document
    engine=pyttsx3.init()
    engine.say("okay! Creating a new document for you!")
    engine.runAndWait()
    
    document = Document()
    document.save('Demo.docx')
    os.system('Demo.docx')
    return "Okay! Creating a new Document for you!"

def launchWord():#Launching word
    engine=pyttsx3.init()
    engine.say("okay! Launching word!")
    engine.runAndWait()
    hi = win32.gencache.EnsureDispatch('Word.Application')
    hi.Visible = True
    return "Okay! Launching Word"

def launchPpt():#Launching powerpoint
    engine=pyttsx3.init()
    engine.say("okay! Launching MS powerpoint!")
    engine.runAndWait()
    gui.press("win")
    gui.typewrite("powerpoint")
    gui.press("enter")
    return "Okay! Launching Powerpoint"

def createPpt():#Creating new powerpoint
    engine=pyttsx3.init()
    engine.say("okay! Creating a new Powepoint Presentation for you!")
    engine.runAndWait()
    import os
    prs = Presentation()
    prs.save("Demo.pptx")
    os.system("Demo.pptx")
    return "Okay! Creating a new Powepoint Presentation for you"

def createExcelSheet():#creating new worksheet
    engine=pyttsx3.init()
    engine.say("okay! Creating a new excel sheet for you!")
    engine.runAndWait()
    import os
    
    wb = Workbook()  
    sheet1 = wb.add_sheet('Sheet 1') 
    wb.save('example.xls')
    os.system("example.xls")
    return "Okay! Creating a new Excel Sheet for you!"

def launchExcel():#Launching Excel
    engine=pyttsx3.init()
    engine.say("okay! Launching Excel")
    engine.runAndWait()
    import pyautogui as gui
    gui.press("win")
    gui.typewrite("excel")
    gui.press("enter")
    return "Okay! Launching Excel"

def takeSelfie():
    #Taking a selife
    import os
    engine=pyttsx3.init()
    engine.say("okay! taking a selfie ...Smile!")
    engine.runAndWait()
    cap=cv2.VideoCapture(0)
    ret,frame =cap.read()
    #gray=cv2.cvtColor(frame,cv2.COLOR_BGR2HSV)
    cv2.imwrite("filename.jpg",frame)
    os.system("filename.jpg")
    return "Okay! Taking a selfie ...Smile!"
def converttorupees():
    engine=pyttsx3.init()
    engine.say("One USD in INR is 69")
    engine.runAndWait()
    conv="One USD in INR is 69"
    return conv
def windows(com):
	command = com
	while True:
	    command = command.upper()
	    
	    if command=="WHAT IS HAPPENING IN THE NBA" or command=="SHOW ME NBA SCORES":
	        return scrapebrowser(command)
	        break
	    elif command == 'HOW IS THE AVENGERS INFINITY WAR' or command == 'HOW GOOD IS THE AVENGERS INFINITY WAR' or command == 'RATINGS OF THE AVENGERS INFINITY WAR':
	        return Avengers(command)
	        break
	    elif command=='WHAT\'S THE TIME' or command=='SHOW ME THE TIME' or command=='WHAT TIME IS IT':
	        return currenttime()
	        break
	    elif command=="WHAT\'S 1 USD IN INR" or command=="CONVERT 1 USD TO INR":
	        return converttorupees()
	        break
	    elif command=="SHUTDOWN" or command=="SHUTDOWN THE COMPUTER":
	        return shutdown()
	        break
	    elif command=="RESTART" or command=="RESTART THE COMPUTER":
	        return restart()
	        break
	    elif command=="SWITCH OFF MONITOR" or command=="GO TO SLEEP":
	        return sleep()
	        break
	    
	    elif command=="GO TO DOWNLOADS" or command=="NAVIGATE ME TO DOWNLOADS" or command=="OPEN DOWNLOADS":
	        return openDownloads()
	        break
	    elif command=="WHAT\'S 9 INTO 2 + 3":
	        engine=pyttsx3.init()
	        engine.say("9 into 2 plus 3 is 21")
	        engine.runAndWait()
	        return "9 into 2 plus 3 is 21" #NOT WORKING
	        break
	    elif command=="CALCULATE 16 TIMES THE SQAURE ROOT OF 29":
	        engine=pyttsx3.init()
	        engine.say("16 TIMES THE SQAURE ROOT OF 29 IS 86.162") #NOT WORKING
	        engine.runAndWait()
	        return "16 times the square root of 29 is 86.162"
	        break
	    elif command=="OPEN CALCULATOR":
	        return calculator()
	        break
	    elif command=="SEND AN EMAIL TO VARUN" or command=="EMAIL VARUN":
	        return sendMail()
	        break
	    elif command=="OPEN COMMAND PROMPT" or command=="GO TO COMMAND PROMPT" or command=="LAUNCH COMMAND PROMPT":
	        return commandPrompt() 
	        break
	    elif command=="CREATE A NEW WORD DOCUMENT":
	        return createDocument()
	        break
	    elif command=="OPEN WORD" or command=="LAUNCH WORD":
	        return launchWord()
	        break
	    elif command=="CREATE A NEW POWERPOINT PRESENTATION":
	        return createPpt()
	        break
	    elif command=="CREATE A NEW EXCEL SHEET":
	        return createExcelSheet()
	        break
	    elif command=="OPEN POWERPOINT" or command=="LAUNCH POWERPOINT":
	        return launchPpt()
	        break
	    elif command=="OPEN EXCEL" or command=="LAUNCH EXCEL":
	        return launchExcel()
	        break
	    elif command=="WHAT\'S THE WEATHER IN BANGALORE" :
	        return weather_info() #not working
	        break
	    elif command=="IS IT GOING TO RAIN":
	        return rain_info()
	        break
	    elif command=="SHOULD I CARRY MY UMBRELLA":
	        return umbrella()
	        break
	    elif command=="TAKE A PICTURE" or command=="TAKE A SELFIE":
	        return takeSelfie()
	        break
	    elif command=="CLOSE THE BROWSER" :
	        return close_browser()
	        break
	    elif command=="CLOSE WORD":
	        return close_word()
	        break
	    elif command=="CLOSE EXCEL":
	         return close_excel()
	         break
	    elif command=="CLOSE POWERPOINT":
	         return close_powerpoint()
	         break
	    elif command=="HOW DO YOU SAY HI HOW ARE YOU IN SPANISH" or command=="HOW DO YOU GREET SOMEONE IN SPANISH":
             return translation()
             break
        

